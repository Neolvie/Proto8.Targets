"""Unit-тесты для парсера DOCX-файлов."""

import io
import pytest
from unittest.mock import patch, MagicMock
from src.services.docx_parser import parse_docx_bytes, parse_docx_file


class TestParseDocxBytes:
    """Тесты функции parse_docx_bytes."""

    def test_invalid_bytes_raises_value_error(self):
        """Невалидные байты вызывают ValueError."""
        with pytest.raises(ValueError, match="Не удалось открыть DOCX"):
            parse_docx_bytes(b"not a docx file")

    def test_empty_bytes_raises_value_error(self):
        """Пустые байты вызывают ValueError."""
        with pytest.raises(ValueError):
            parse_docx_bytes(b"")

    def test_parse_real_docx_if_available(self):
        """Интеграционный тест: парсинг реального Ario.docx если доступен."""
        import os
        path = os.path.join("data", "Ario.docx")
        if not os.path.exists(path):
            pytest.skip("Ario.docx недоступен")
        result = parse_docx_file(path)
        assert isinstance(result, str)
        assert len(result) > 0


class TestParseDocxFile:
    """Тесты функции parse_docx_file."""

    def test_file_not_found_raises(self):
        """Несуществующий файл вызывает FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            parse_docx_file("/nonexistent/path/file.docx")

    def test_invalid_file_raises_value_error(self, tmp_path):
        """Невалидный DOCX-файл вызывает ValueError."""
        bad_file = tmp_path / "bad.docx"
        bad_file.write_bytes(b"this is not a docx")
        with pytest.raises(ValueError, match="Не удалось открыть DOCX"):
            parse_docx_file(str(bad_file))


class TestExtractText:
    """Тесты извлечения текста через mock."""

    def test_extract_paragraphs(self):
        """Абзацы корректно извлекаются из документа."""
        from src.services.docx_parser import _extract_document_text
        from docx import Document

        doc = Document()
        doc.add_paragraph("Первый абзац")
        doc.add_paragraph("Второй абзац")

        result = _extract_document_text(doc)
        assert "Первый абзац" in result
        assert "Второй абзац" in result

    def test_extract_headings(self):
        """Заголовки корректно извлекаются с маркерами уровня."""
        from src.services.docx_parser import _extract_document_text
        from docx import Document

        doc = Document()
        doc.add_heading("Главный заголовок", level=1)
        doc.add_paragraph("Текст под заголовком")

        result = _extract_document_text(doc)
        assert "Главный заголовок" in result
        assert "Текст под заголовком" in result

    def test_extract_table(self):
        """Таблица извлекается в виде текста."""
        from src.services.docx_parser import _extract_document_text
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Ячейка 1"
        table.cell(0, 1).text = "Ячейка 2"
        table.cell(1, 0).text = "Ячейка 3"
        table.cell(1, 1).text = "Ячейка 4"

        result = _extract_document_text(doc)
        assert "Ячейка 1" in result
        assert "ТАБЛИЦА" in result

    def test_skip_empty_paragraphs(self):
        """Пустые абзацы не добавляются в результат."""
        from src.services.docx_parser import _extract_document_text
        from docx import Document

        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("Текст")

        result = _extract_document_text(doc)
        lines = [l for l in result.split("\n") if l.strip()]
        assert len(lines) == 1
        assert "Текст" in lines[0]
